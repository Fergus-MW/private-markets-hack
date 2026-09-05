from tempfile import TemporaryDirectory
from pathlib import Path
import subprocess,httpx
from docx import Document
with TemporaryDirectory() as directory:
    root=Path(directory); doc=Document(); doc.add_heading('Quarterly accounts', 1)
    table=doc.add_table(rows=1, cols=3); table.style='Table Grid'
    for cell,value in zip(table.rows[0].cells,['Metric','2025','2026']):cell.text=value
    for values in [('Revenue','100','120'),('EBITDA','20','25'),('Cash','40','45')]:
        for cell,value in zip(table.add_row().cells, values):cell.text=value
    doc.save(root/'table.docx')
    subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',directory,str(root/'table.docx')],check=True,capture_output=True)
    with httpx.Client(timeout=480) as client:
        r=client.post('http://127.0.0.1:8080/documents?pdf_strategy=hi_res',files={'file':('table.pdf',(root/'table.pdf').read_bytes())})
        assert r.status_code==200,r.text
        elements=client.get('http://127.0.0.1:8080/documents/'+r.json()['document_id']+'/elements').json()['elements']
        assert any(e['metadata'].get('text_as_html') for e in elements),elements
        print('PASS: high-resolution PDF table HTML')
